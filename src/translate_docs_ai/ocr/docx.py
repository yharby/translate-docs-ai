"""
Direct text extraction for DOCX files.

Uses python-docx to extract text without OCR.
"""

from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table

from translate_docs_ai.ocr.base import OCRProvider, OCRQuality, OCRResult


class DocxExtractor(OCRProvider):
    """
    Extract text from DOCX files using python-docx.

    No OCR needed - extracts text directly from the document structure.
    Preserves basic formatting as markdown.
    """

    @property
    def name(self) -> str:
        return "docx_direct"

    def can_handle(self, file_path: Path) -> bool:
        """Check if this extractor can handle the file."""
        return file_path.suffix.lower() in {".docx", ".doc"}

    def _table_to_markdown(self, table: Table) -> str:
        """Convert a DOCX table to markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) >= 1:
            # Add header separator after first row
            num_cols = len(table.rows[0].cells) if table.rows else 0
            separator = "| " + " | ".join(["---"] * num_cols) + " |"
            rows.insert(1, separator)

        return "\n".join(rows)

    def _extract_content_as_markdown(self, doc: Document) -> str:
        """
        Extract document content and convert to markdown format.

        Handles paragraphs, headings, lists, and tables.
        """
        content_parts = []

        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith("p"):
                # Find the corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        text = para.text.strip()
                        if not text:
                            continue

                        # Check for heading styles
                        style_name = para.style.name if para.style else ""
                        if style_name.startswith("Heading"):
                            try:
                                level = int(
                                    style_name.replace("Heading ", "").replace("Heading", "1")
                                )
                                level = min(level, 6)  # Max 6 levels in markdown
                            except ValueError:
                                level = 1
                            content_parts.append(f"{'#' * level} {text}")
                        elif style_name == "Title":
                            content_parts.append(f"# {text}")
                        elif style_name == "Subtitle":
                            content_parts.append(f"## {text}")
                        elif "List" in style_name or para.style.name.startswith("List"):
                            # Handle list items
                            content_parts.append(f"- {text}")
                        else:
                            content_parts.append(text)
                        break

            # Handle tables
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element is element:
                        content_parts.append("")  # Empty line before table
                        content_parts.append(self._table_to_markdown(table))
                        content_parts.append("")  # Empty line after table
                        break

        return "\n\n".join(content_parts)

    async def extract_page(
        self,
        file_path: Path,
        page_number: int,
        **kwargs: Any,
    ) -> OCRResult:
        """
        Extract text from a DOCX file.

        DOCX files don't have traditional pages, so we treat the entire
        document as a single "page" (page 0).

        Args:
            file_path: Path to DOCX file.
            page_number: Page number (0-indexed). Only page 0 has content.
            **kwargs: Ignored for DOCX files.

        Returns:
            OCRResult with document content (for page 0) or empty (for other pages).
        """
        if page_number != 0:
            # DOCX files are treated as single page
            return OCRResult(
                content="",
                page_number=page_number,
                confidence=1.0,
                quality=OCRQuality.EXCELLENT,
                model_used=self.name,
                metadata={"note": "DOCX files are treated as a single page"},
            )

        try:
            doc = Document(str(file_path))
            content = self._extract_content_as_markdown(doc)

            # Get document metadata
            core_props = doc.core_properties
            metadata = {
                "file_type": ".docx",
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            return OCRResult(
                content=content,
                page_number=0,
                confidence=1.0,  # Direct extraction = perfect confidence
                quality=OCRQuality.EXCELLENT,
                model_used=self.name,
                metadata=metadata,
            )

        except PackageNotFoundError:
            return OCRResult(
                content="",
                page_number=0,
                confidence=0.0,
                quality=OCRQuality.POOR,
                model_used=self.name,
                metadata={"error": "Invalid or corrupted DOCX file"},
            )

        except Exception as e:
            return OCRResult(
                content="",
                page_number=0,
                confidence=0.0,
                quality=OCRQuality.POOR,
                model_used=self.name,
                metadata={"error": str(e)},
            )

    async def extract_document(
        self,
        file_path: Path,
        **kwargs: Any,
    ) -> list[OCRResult]:
        """
        Extract text from a DOCX file.

        Returns a single-element list since DOCX files are treated as one page.

        Args:
            file_path: Path to DOCX file.
            **kwargs: Ignored for DOCX files.

        Returns:
            List with one OCRResult containing the document content.
        """
        result = await self.extract_page(file_path, 0, **kwargs)
        return [result]

    def get_page_count(self, file_path: Path) -> int:
        """
        Get the number of pages in a DOCX file.

        DOCX files are treated as a single page for processing purposes.

        Args:
            file_path: Path to DOCX file.

        Returns:
            Always returns 1.
        """
        return 1
