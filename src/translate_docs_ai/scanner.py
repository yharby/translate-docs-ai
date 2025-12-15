"""
Document scanner for translate-docs-ai.

Recursively scans directories for documents and catalogs them in DuckDB.
"""

from __future__ import annotations

import mimetypes
from collections.abc import Iterator
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn

from translate_docs_ai.database import Database, Document, Status

# Supported file extensions and their types
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
}


class DocumentScanner:
    """Scans directories for documents and catalogs them."""

    def __init__(self, db: Database, console: Console | None = None):
        """
        Initialize scanner.

        Args:
            db: Database instance for storing document catalog.
            console: Rich console for output. If None, creates a new one.
        """
        self.db = db
        self.console = console or Console()

    def scan_directory(
        self,
        input_dir: Path,
        recursive: bool = True,
    ) -> list[Document]:
        """
        Scan a directory for documents and add them to the catalog.

        Args:
            input_dir: Directory to scan.
            recursive: Whether to scan subdirectories.

        Returns:
            List of Document objects that were added.
        """
        input_dir = Path(input_dir).resolve()
        if not input_dir.exists():
            raise FileNotFoundError(f"Input directory not found: {input_dir}")

        if not input_dir.is_dir():
            raise ValueError(f"Not a directory: {input_dir}")

        documents: list[Document] = []
        files = list(self._find_documents(input_dir, recursive))

        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]{task.description}"),
            console=self.console,
        ) as progress:
            task = progress.add_task(f"Scanning {input_dir.name}...", total=len(files))

            for file_path in files:
                progress.update(task, description=f"Processing {file_path.name}")

                # Check if document already exists
                existing = self.db.get_document_by_path(str(file_path))
                if existing:
                    documents.append(existing)
                    progress.advance(task)
                    continue

                # Get document info
                doc = self._create_document_record(file_path, input_dir)
                if doc:
                    doc_id = self.db.add_document(doc)
                    doc.id = doc_id
                    documents.append(doc)

                    # Log the addition
                    self.db.log(
                        level="INFO",
                        stage="scan",
                        message=f"Added document: {doc.file_name}",
                        context={"path": str(file_path), "pages": doc.total_pages},
                    )

                progress.advance(task)

        self.console.print(f"[green]Found {len(documents)} documents[/green]")
        return documents

    def _find_documents(self, directory: Path, recursive: bool = True) -> Iterator[Path]:
        """Find all supported documents in a directory."""
        pattern = "**/*" if recursive else "*"

        for path in directory.glob(pattern):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                yield path

    def _create_document_record(self, file_path: Path, base_dir: Path) -> Document | None:
        """Create a Document record from a file path."""
        try:
            file_type = SUPPORTED_EXTENSIONS.get(file_path.suffix.lower())
            if not file_type:
                return None

            relative_path = file_path.relative_to(base_dir)
            total_pages = self._count_pages(file_path, file_type)

            return Document(
                file_name=file_path.name,
                full_path=str(file_path),
                relative_path=str(relative_path),
                file_type=file_type,
                total_pages=total_pages,
                status=Status.PENDING,
            )
        except Exception as e:
            self.console.print(f"[yellow]Warning: Could not process {file_path}: {e}[/yellow]")
            self.db.log(
                level="WARNING",
                stage="scan",
                message=f"Could not process file: {e}",
                context={"path": str(file_path)},
            )
            return None

    def _count_pages(self, file_path: Path, file_type: str) -> int:
        """Count pages in a document."""
        try:
            if file_type == "pdf":
                with fitz.open(file_path) as doc:
                    return len(doc)
            elif file_type == "docx":
                # DOCX doesn't have a direct page count, estimate from sections
                doc = DocxDocument(file_path)
                # Rough estimate: ~500 words per page
                word_count = sum(len(p.text.split()) for p in doc.paragraphs)
                return max(1, word_count // 500)
            elif file_type in ("markdown", "text"):
                # Count lines and estimate pages (~50 lines per page)
                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    lines = sum(1 for _ in f)
                return max(1, lines // 50)
            else:
                return 1
        except Exception:
            return 1

    def get_document_status(self) -> dict[str, int]:
        """Get counts of documents by status."""
        docs = self.db.get_all_documents()
        status_counts: dict[str, int] = {}
        for doc in docs:
            status = doc.status.value
            status_counts[status] = status_counts.get(status, 0) + 1
        return status_counts

    def get_pending_documents(self) -> list[Document]:
        """Get all documents that haven't been processed yet."""
        return self.db.get_all_documents(Status.PENDING)

    def rescan_directory(self, input_dir: Path) -> list[Document]:
        """
        Rescan a directory, adding new documents and skipping existing ones.

        This is useful for incremental updates when new files are added.
        """
        return self.scan_directory(input_dir, recursive=True)


def get_mime_type(file_path: Path) -> str | None:
    """Get MIME type for a file."""
    mime_type, _ = mimetypes.guess_type(str(file_path))
    return mime_type


def is_supported_file(file_path: Path) -> bool:
    """Check if a file is supported for processing."""
    return file_path.suffix.lower() in SUPPORTED_EXTENSIONS
