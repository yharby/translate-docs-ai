"""
DOCX exporter for translated documents.

Uses python-docx for professional Word document generation.
Uses mistune for proper markdown AST parsing.
Handles RTL to LTR table column reversal for proper reading direction.
Supports applying extracted styling from source PDF documents.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Any

import mistune
from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from translate_docs_ai.export.table_utils import process_content_tables
from translate_docs_ai.styling import FontManager, get_fallback_font

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

    from translate_docs_ai.database import Database, Document


@dataclass
class DOCXExportResult:
    """Result of a DOCX export operation."""

    document_id: int
    document_name: str
    output_path: Path
    pages_exported: int
    language: str
    success: bool
    error: str | None = None


class DocxRenderer(mistune.BaseRenderer):
    """
    Custom mistune renderer that outputs to python-docx Document.

    Instead of returning strings (like HTML), this renderer directly
    adds elements to a Word document.
    """

    # Color scheme
    COLORS = {
        "primary": RGBColor(0x1A, 0x1A, 0x2E),  # Dark blue
        "secondary": RGBColor(0x4A, 0x90, 0xD9),  # Light blue
        "text": RGBColor(0x33, 0x33, 0x33),  # Dark gray
        "light": RGBColor(0x66, 0x66, 0x66),  # Medium gray
    }

    def __init__(self, docx: DocxDocument, is_rtl: bool = False) -> None:
        super().__init__()
        self.docx = docx
        self.is_rtl = is_rtl
        self._current_paragraph: Paragraph | None = None
        self._list_level = 0

    def _align_paragraph(self, para: Paragraph) -> None:
        """Apply RTL alignment if needed."""
        if self.is_rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Block-level elements

    def paragraph(self, text: str) -> str:
        """Handle paragraph - text already contains rendered inline content."""
        if text.strip():
            para = self.docx.add_paragraph()
            self._current_paragraph = para
            self._add_inline_content(para, text)
            self._align_paragraph(para)
        return ""

    def heading(self, text: str, level: int, **attrs: Any) -> str:
        """Handle headings H1-H6."""
        # Strip any HTML tags that might be in the text
        clean_text = re.sub(r"<[^>]+>", "", text)
        h = self.docx.add_heading(clean_text, level=level)
        self._align_paragraph(h)
        return ""

    def thematic_break(self) -> str:
        """Handle horizontal rules (---, ***, ___)."""
        para = self.docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("─" * 40)
        run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
        return ""

    def block_quote(self, text: str) -> str:
        """Handle blockquotes."""
        para = self.docx.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        if self.is_rtl:
            para.paragraph_format.right_indent = Inches(0.5)
            para.paragraph_format.left_indent = Inches(0)
        run = para.add_run(re.sub(r"<[^>]+>", "", text).strip())
        run.font.italic = True
        run.font.color.rgb = self.COLORS["light"]
        self._align_paragraph(para)
        return ""

    def block_code(self, code: str, info: str | None = None) -> str:
        """Handle fenced code blocks."""
        para = self.docx.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)

        # Add dark background shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2D2D2D"/>')
        para._p.get_or_add_pPr().append(shading)

        run = para.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xF8, 0xF8, 0xF2)
        return ""

    def list(self, text: str, ordered: bool, **attrs: Any) -> str:
        """Handle lists - text contains rendered list items."""
        # List items are already added by list_item, just return empty
        return ""

    def list_item(self, text: str, **attrs: Any) -> str:
        """Handle list items."""
        # Clean up the text - remove nested paragraph tags
        clean_text = re.sub(r"<[^>]+>", "", text).strip()

        # Determine style based on context (we track this via attrs or pattern)
        # Default to bullet, but check if parent is ordered
        style = "List Bullet"

        para = self.docx.add_paragraph(clean_text, style=style)
        self._align_paragraph(para)
        return ""

    def table(self, text: str) -> str:
        """Handle tables - parse the accumulated rows."""
        # Tables are handled specially - we accumulate rows and render at once
        return ""

    # Inline elements - these return marked-up text for paragraph processing

    def text(self, text: str) -> str:
        """Plain text."""
        return text

    def emphasis(self, text: str) -> str:
        """Italic text."""
        return f"<em>{text}</em>"

    def strong(self, text: str) -> str:
        """Bold text."""
        return f"<strong>{text}</strong>"

    def codespan(self, text: str) -> str:
        """Inline code."""
        return f"<code>{text}</code>"

    def linebreak(self) -> str:
        """Line break."""
        return "\n"

    def softbreak(self) -> str:
        """Soft line break."""
        return " "

    def link(self, text: str, url: str, title: str | None = None) -> str:
        """Handle links - just return the text for now."""
        return text

    def image(self, text: str, url: str, title: str | None = None) -> str:
        """Handle images - just return alt text."""
        return f"[{text}]"

    def _add_inline_content(self, para: Paragraph, text: str) -> None:
        """Parse inline formatting and add runs to paragraph."""
        # Pattern to find our inline markers
        pattern = r"(<strong>.*?</strong>|<em>.*?</em>|<code>.*?</code>)"

        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith("<strong>") and part.endswith("</strong>"):
                content = part[8:-9]  # Remove <strong></strong>
                run = para.add_run(content)
                run.font.bold = True
            elif part.startswith("<em>") and part.endswith("</em>"):
                content = part[4:-5]  # Remove <em></em>
                run = para.add_run(content)
                run.font.italic = True
            elif part.startswith("<code>") and part.endswith("</code>"):
                content = part[6:-7]  # Remove <code></code>
                run = para.add_run(content)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            else:
                para.add_run(part)


class DOCXExporter:
    """
    Exports translated documents to DOCX files.

    Creates professionally formatted Word documents with:
    - Custom styling
    - RTL support for Arabic
    - Tables with borders
    - Proper markdown parsing via mistune
    """

    LANGUAGE_COLUMNS = {
        "en": "en_content",
        "ar": "ar_content",
        "fr": "fr_content",
    }

    RTL_LANGUAGES = {"ar", "he", "fa", "ur"}

    # Color scheme
    COLORS = {
        "primary": RGBColor(0x1A, 0x1A, 0x2E),  # Dark blue
        "secondary": RGBColor(0x4A, 0x90, 0xD9),  # Light blue
        "text": RGBColor(0x33, 0x33, 0x33),  # Dark gray
        "light": RGBColor(0x66, 0x66, 0x66),  # Medium gray
    }

    def __init__(self, db: Database, output_dir: Path) -> None:
        """
        Initialize the DOCX exporter.

        Args:
            db: Database instance.
            output_dir: Base output directory.
        """
        self.db = db
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._font_manager = FontManager()
        self._doc_style: dict[str, Any] | None = None

    def _load_document_styling(self, document_id: int) -> dict[str, Any] | None:
        """Load extracted styling metadata for a document."""
        return self.db.get_document_styling(document_id)

    def _get_font_for_language(self, language: str, original_font: str | None = None) -> str:
        """Get appropriate font for language with fallback support."""
        if original_font:
            return get_fallback_font(original_font, language)
        return get_fallback_font("Arial", language)

    def _hex_to_rgb(self, hex_color: str | None) -> RGBColor | None:
        """Convert hex color string to RGBColor."""
        if not hex_color or not hex_color.startswith("#"):
            return None
        try:
            hex_color = hex_color.lstrip("#")
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        except (ValueError, IndexError):
            return None

    def export_document(
        self,
        document: Document,
        language: str = "en",
        source_lang: str | None = None,
        clean: bool = False,
    ) -> DOCXExportResult:
        """
        Export a single document to DOCX.

        Args:
            document: Document to export.
            language: Target language code (en, ar, fr).
            source_lang: Source language code for RTL/LTR table conversion.
            clean: If True, export without metadata headers (currently unused for DOCX).

        Returns:
            DOCXExportResult with export status.
        """
        # Note: clean parameter accepted for API consistency but DOCX export
        # doesn't add metadata headers like Markdown/PDF exports do
        _ = clean  # Unused for now, here for API consistency
        content_column = self.LANGUAGE_COLUMNS.get(language)
        if not content_column:
            return DOCXExportResult(
                document_id=document.id,
                document_name=document.file_name,
                output_path=self.output_dir,
                pages_exported=0,
                language=language,
                success=False,
                error=f"Unsupported language: {language}",
            )

        pages = self.db.get_document_pages(document.id)

        translated_pages: list[tuple[int, str]] = []
        for page in pages:
            content = getattr(page, content_column, None)
            if content:
                # Process tables: convert HTML to MD and reverse columns for RTL→LTR
                if source_lang:
                    content = process_content_tables(content, source_lang, language)
                translated_pages.append((page.page_number, content))

        if not translated_pages:
            return DOCXExportResult(
                document_id=document.id,
                document_name=document.file_name,
                output_path=self.output_dir,
                pages_exported=0,
                language=language,
                success=False,
                error=f"No {language} translations found",
            )

        translated_pages.sort(key=lambda x: x[0])

        # Create safe filename
        doc_stem = Path(document.file_name).stem
        safe_name = self._sanitize_filename(doc_stem)
        output_file = self.output_dir / f"{safe_name}_{language}.docx"

        # Load extracted styling from source document
        self._doc_style = self._load_document_styling(document.id)

        try:
            # Create document
            docx = DocxDocument()
            is_rtl = language in self.RTL_LANGUAGES

            # Setup styles with extracted styling info
            self._setup_styles(docx, is_rtl, language)

            # Combine all content for continuous flow
            all_content = "\n\n".join(content for _, content in translated_pages)

            # Pre-process: normalize Unicode bullets to markdown
            all_content = self._normalize_bullets(all_content)

            # Parse and render markdown using mistune with custom renderer
            self._render_markdown_to_docx(docx, all_content, is_rtl)

            # Save
            docx.save(str(output_file))

            return DOCXExportResult(
                document_id=document.id,
                document_name=document.file_name,
                output_path=output_file,
                pages_exported=len(translated_pages),
                language=language,
                success=True,
            )

        except Exception as e:
            return DOCXExportResult(
                document_id=document.id,
                document_name=document.file_name,
                output_path=self.output_dir,
                pages_exported=0,
                language=language,
                success=False,
                error=str(e),
            )

    def _normalize_bullets(self, content: str) -> str:
        """Convert Unicode bullet characters to markdown list syntax."""
        lines = content.split("\n")
        result = []

        for line in lines:
            stripped = line.lstrip()
            indent = len(line) - len(stripped)

            # Convert Unicode bullets to markdown
            # • (U+2022), ● (U+25CF), ○ (U+25CB), ▪ (U+25AA), ▫ (U+25AB), ◦ (U+25E6)
            if stripped.startswith(("• ", "● ", "○ ", "▪ ", "▫ ", "◦ ")):
                indent_str = " " * indent
                bullet_text = stripped[2:].strip()
                result.append(f"{indent_str}- {bullet_text}")
            elif stripped.startswith(("•", "●", "○", "▪", "▫", "◦")) and len(stripped) > 1:
                # Bullets without space after
                indent_str = " " * indent
                bullet_text = stripped[1:].strip()
                result.append(f"{indent_str}- {bullet_text}")
            else:
                result.append(line)

        return "\n".join(result)

    def _render_markdown_to_docx(
        self,
        docx: DocxDocument,
        content: str,
        is_rtl: bool,
    ) -> None:
        """
        Parse markdown and render to DOCX using hybrid approach.

        Uses mistune for proper AST parsing, then renders each token
        appropriately to the Word document.
        """
        # Create markdown parser with table plugin, renderer=None returns AST
        md = mistune.create_markdown(renderer=None, plugins=["table", "strikethrough"])

        # Parse to AST - returns (tokens, state) tuple
        result = md.parse(content)
        tokens = result[0] if isinstance(result, tuple) else result

        # Render each token to the document
        self._render_tokens(docx, tokens, is_rtl)

    def _render_tokens(
        self,
        docx: DocxDocument,
        tokens: list[dict[str, Any]],
        is_rtl: bool,
    ) -> None:
        """Render AST tokens to Word document."""
        for token in tokens:
            token_type = token.get("type")

            if token_type == "heading":
                level = token.get("attrs", {}).get("level", 1)
                text = self._extract_text(token.get("children", []))
                h = docx.add_heading(text, level=level)
                if is_rtl:
                    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            elif token_type == "paragraph":
                children = token.get("children", [])
                if children:
                    para = docx.add_paragraph()
                    self._render_inline(para, children)
                    if is_rtl:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            elif token_type == "thematic_break":
                para = docx.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run("─" * 40)
                run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

            elif token_type == "block_code":
                code = token.get("raw", "")
                self._add_code_block(docx, code)

            elif token_type == "list":
                ordered = token.get("attrs", {}).get("ordered", False)
                items = token.get("children", [])
                self._render_list(docx, items, ordered, is_rtl)

            elif token_type == "block_quote":
                children = token.get("children", [])
                text = self._extract_text(children)
                para = docx.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(text)
                run.font.italic = True
                run.font.color.rgb = self.COLORS["light"]
                if is_rtl:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            elif token_type == "table":
                self._render_table(docx, token, is_rtl)

            elif token_type == "blank_line":
                # Skip blank lines or add minimal spacing
                pass

    def _render_list(
        self,
        docx: DocxDocument,
        items: list[dict[str, Any]],
        ordered: bool,
        is_rtl: bool,
        level: int = 0,
    ) -> None:
        """Render a list with proper nesting support."""
        style = "List Number" if ordered else "List Bullet"

        # Use different styles for nested levels
        if level > 0:
            style = f"{style} {min(level + 1, 5)}"
            # Check if style exists, fall back to base style
            try:
                docx.styles[style]
            except KeyError:
                style = "List Number" if ordered else "List Bullet"

        for item in items:
            if item.get("type") == "list_item":
                children = item.get("children", [])
                # Extract text content and any nested lists
                text_parts = []
                nested_lists = []

                for child in children:
                    if child.get("type") == "list":
                        nested_lists.append(child)
                    elif child.get("type") == "paragraph":
                        text_parts.append(self._extract_text(child.get("children", [])))
                    else:
                        text_parts.append(self._extract_text([child]))

                text = " ".join(text_parts).strip()
                if text:
                    para = docx.add_paragraph(style=style)
                    self._render_inline(para, [{"type": "text", "raw": text}])
                    if is_rtl:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Handle nested lists
                for nested in nested_lists:
                    nested_ordered = nested.get("attrs", {}).get("ordered", False)
                    self._render_list(
                        docx,
                        nested.get("children", []),
                        nested_ordered,
                        is_rtl,
                        level + 1,
                    )

    def _render_table(
        self,
        docx: DocxDocument,
        token: dict[str, Any],
        is_rtl: bool,
    ) -> None:
        """Render a markdown table to Word table."""
        children = token.get("children", [])
        if not children:
            return

        # Extract rows - mistune table structure:
        # table -> table_head (contains cells directly) + table_body (contains rows)
        rows_data: list[list[str]] = []
        for child in children:
            child_type = child.get("type")

            if child_type == "table_head":
                # Header cells are direct children of table_head
                cells = []
                for cell in child.get("children", []):
                    if cell.get("type") == "table_cell":
                        cell_text = self._extract_text(cell.get("children", []))
                        cells.append(cell_text)
                if cells:
                    rows_data.append(cells)

            elif child_type == "table_body":
                # Body contains rows, each row contains cells
                for row in child.get("children", []):
                    if row.get("type") == "table_row":
                        cells = []
                        for cell in row.get("children", []):
                            if cell.get("type") == "table_cell":
                                cell_text = self._extract_text(cell.get("children", []))
                                cells.append(cell_text)
                        if cells:
                            rows_data.append(cells)

        if not rows_data:
            return

        # Create table
        num_cols = max(len(row) for row in rows_data)
        table = docx.add_table(rows=len(rows_data), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT if is_rtl else WD_TABLE_ALIGNMENT.LEFT

        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    run = para.add_run(cell_text)

                    run.font.size = Pt(8)
                    run.font.name = "Segoe UI"

                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)

                    # Header row styling (first row)
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4A90D9"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
                    elif i % 2 == 0:
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                        cell._tc.get_or_add_tcPr().append(shading)

        docx.add_paragraph()  # Space after table

    def _render_inline(
        self,
        para: Paragraph,
        children: list[dict[str, Any]],
    ) -> None:
        """Render inline content (text, bold, italic, code, links)."""
        for child in children:
            child_type = child.get("type")

            if child_type == "text":
                para.add_run(child.get("raw", ""))

            elif child_type == "strong":
                text = self._extract_text(child.get("children", []))
                run = para.add_run(text)
                run.font.bold = True

            elif child_type == "emphasis":
                text = self._extract_text(child.get("children", []))
                run = para.add_run(text)
                run.font.italic = True

            elif child_type == "codespan":
                text = child.get("raw", "")
                run = para.add_run(text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)

            elif child_type == "link":
                text = self._extract_text(child.get("children", []))
                para.add_run(text)

            elif child_type == "softbreak":
                para.add_run(" ")

            elif child_type == "linebreak":
                para.add_run("\n")

            else:
                # Fallback: try to extract text
                text = self._extract_text([child])
                if text:
                    para.add_run(text)

    def _extract_text(self, children: list[dict[str, Any]]) -> str:
        """Extract plain text from nested token structure."""
        parts = []
        for child in children:
            if child.get("type") == "text":
                parts.append(child.get("raw", ""))
            elif "children" in child:
                parts.append(self._extract_text(child["children"]))
            elif "raw" in child:
                parts.append(child["raw"])
        return "".join(parts)

    def _add_code_block(self, docx: DocxDocument, code: str) -> None:
        """Add a formatted code block."""
        para = docx.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)

        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2D2D2D"/>')
        para._p.get_or_add_pPr().append(shading)

        run = para.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xF8, 0xF8, 0xF2)

    def _setup_styles(self, docx: DocxDocument, is_rtl: bool, language: str = "en") -> None:
        """Setup custom styles for the document."""
        styles = docx.styles

        # Fixed sizes for consistent, readable output
        body_size = 10.0  # 10pt for body text
        body_color = self.COLORS["text"]

        # Get appropriate fonts for the target language
        body_font = None
        heading_font = None
        if self._doc_style:
            body_font = self._doc_style.get("dominant_font")
            heading_font = self._doc_style.get("heading_font") or body_font
            if self._doc_style.get("dominant_color"):
                extracted_color = self._hex_to_rgb(self._doc_style["dominant_color"])
                if extracted_color:
                    body_color = extracted_color

        body_font_name = self._get_font_for_language(language, body_font)
        heading_font_name = self._get_font_for_language(language, heading_font)

        # Title style
        if "CustomTitle" not in [s.name for s in styles]:
            title_style = styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.font.name = heading_font_name
            title_style.font.color.rgb = self.COLORS["primary"]
            title_style.paragraph_format.space_after = Pt(12)

        # Heading styles - clear hierarchy with good size differentiation
        heading_sizes = {1: 14, 2: 12, 3: 11, 4: 10.5, 5: 10, 6: 10}
        for level in range(1, 7):
            style_name = f"Heading {level}"
            h_style = styles[style_name]
            h_style.font.size = Pt(heading_sizes[level])
            h_style.font.bold = True
            h_style.font.name = heading_font_name

            if level == 1:
                h_style.font.color.rgb = self.COLORS["primary"]
                h_style.paragraph_format.space_before = Pt(12)
                h_style.paragraph_format.space_after = Pt(6)
            elif level == 2:
                h_style.font.color.rgb = self.COLORS["secondary"]
                h_style.paragraph_format.space_before = Pt(10)
                h_style.paragraph_format.space_after = Pt(4)
            else:
                h_style.font.color.rgb = self.COLORS["text"]
                h_style.paragraph_format.space_before = Pt(8)
                h_style.paragraph_format.space_after = Pt(3)

        # Normal text - 10pt for readability
        normal_style = styles["Normal"]
        normal_style.font.size = Pt(body_size)
        normal_style.font.name = body_font_name
        normal_style.font.color.rgb = body_color
        normal_style.paragraph_format.space_after = Pt(4)
        normal_style.paragraph_format.space_before = Pt(2)
        normal_style.paragraph_format.line_spacing = 1.15

        # List styles - same size as body text
        for style_name in ["List Bullet", "List Number"]:
            if style_name in [s.name for s in styles]:
                list_style = styles[style_name]
                list_style.font.size = Pt(body_size)
                list_style.font.name = body_font_name
                list_style.paragraph_format.space_after = Pt(2)

        if is_rtl:
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize filename for filesystem compatibility."""
        return "".join(c if c.isalnum() or c in "._- " else "_" for c in name)

    def export_all(
        self,
        documents: list[Document],
        language: str = "en",
        source_lang: str | None = None,
    ) -> list[DOCXExportResult]:
        """Export multiple documents to DOCX."""
        results = []
        for doc in documents:
            result = self.export_document(doc, language, source_lang)
            results.append(result)
        return results
